#!/usr/bin/env python3
"""Run readiness checks and a safe local audit for migrated skills."""

from __future__ import annotations

import argparse
import importlib.util
import json
import os
import shutil
import subprocess
import sys
from pathlib import Path
from typing import Any, Optional

if __package__ is None or __package__ == "":
    sys.path.insert(0, str(Path(__file__).resolve().parent))

from runtime_registry import NOTION_SKILLS, load_or_probe_registry, registry_path, resolve_runtime_root  # type: ignore  # noqa: E402
from skill_inventory import discover_skills  # type: ignore  # noqa: E402
from utils import dump_json, dump_text, ensure_lab_workspace, load_json, utc_now_iso  # type: ignore  # noqa: E402
from validate_skill import validate_skill  # type: ignore  # noqa: E402

DEFAULT_SKILLS_ROOT = Path.home() / ".codex" / "skills"
SAFE_BENCHMARK_SKILLS = {
    "doc",
    "pdf",
    "jupyter-notebook",
    "latex-to-word",
    "openai-docs",
    "playwright",
    "screenshot",
    "gh-fix-ci",
}
READY_EXTERNAL_WRITE_SKILLS = set(NOTION_SKILLS) | {"yeet"}


def has_python_module(module_name: str) -> bool:
    return importlib.util.find_spec(module_name) is not None


def command_exists(command: str) -> bool:
    return shutil.which(command) is not None


def run_command(command: list[str], *, timeout: int = 20) -> tuple[int, str, str]:
    completed = subprocess.run(command, capture_output=True, text=True, timeout=timeout, check=False)
    return completed.returncode, completed.stdout.strip(), completed.stderr.strip()


def probe_codex_mcp() -> dict[str, Any]:
    if not command_exists("codex"):
        return {
            "codex_present": False,
            "configured": False,
            "servers": [],
            "stdout": "",
            "stderr": "codex missing",
            "return_code": 127,
        }
    rc, stdout, stderr = run_command(
        ["codex", "-c", 'model_reasoning_effort="medium"', "mcp", "list"]
    )
    lines = [line.strip() for line in stdout.splitlines() if line.strip()]
    configured = rc == 0 and any(not line.lower().startswith("no mcp servers") for line in lines)
    return {
        "codex_present": True,
        "configured": configured,
        "servers": lines,
        "stdout": stdout,
        "stderr": stderr,
        "return_code": rc,
    }


def probe_environment(*, runtime_root: str | Path | None = None) -> dict[str, Any]:
    runtime_registry = load_or_probe_registry(runtime_root=runtime_root, install_missing=False)
    core_probe = runtime_registry.get("domains", {}).get("core-tools", {}).get("probe", {})
    docs_probe = runtime_registry.get("domains", {}).get("docs-python", {}).get("probe", {})

    config_path = Path.home() / ".codex" / "config.toml"
    config_text = config_path.read_text(encoding="utf-8") if config_path.exists() else ""
    mcp = probe_codex_mcp()
    notion_configured = any("notion" in line.lower() for line in mcp.get("servers", [])) or bool(core_probe.get("notion_mcp_configured")) or "notion" in config_text.lower()
    playwright_mcp_configured = any("playwright" in line.lower() for line in mcp.get("servers", [])) or bool(core_probe.get("playwright_mcp_configured")) or "playwright" in config_text.lower()
    rmcp_enabled = bool(core_probe.get("rmcp_enabled")) or "rmcp_client" in config_text.lower() or "experimental_use_rmcp_client" in config_text.lower()

    return {
        "python3_present": command_exists("python3"),
        "git_present": command_exists("git"),
        "pandoc_present": bool(core_probe.get("pandoc_present")) or command_exists("pandoc"),
        "playwright_present": bool(core_probe.get("playwright_present")) or command_exists("playwright"),
        "playwright_browser_cache_present": bool(core_probe.get("playwright_browser_cache_present")),
        "playwright_mcp_configured": playwright_mcp_configured,
        "gh_present": bool(core_probe.get("gh_present")) or command_exists("gh"),
        "gh_authenticated": bool(core_probe.get("gh_authenticated")),
        "gh_auth_message": core_probe.get("gh_auth_message", "gh missing"),
        "display": os.environ.get("DISPLAY"),
        "wayland_display": os.environ.get("WAYLAND_DISPLAY"),
        "linux_screenshot_tools": core_probe.get(
            "linux_screenshot_tools",
            {
                "scrot": command_exists("scrot"),
                "gnome-screenshot": command_exists("gnome-screenshot"),
                "import": command_exists("import"),
            },
        ),
        "python_modules": {
            "docx": bool(docs_probe.get("modules", {}).get("docx")) or has_python_module("docx"),
            "pdf2image": bool(docs_probe.get("modules", {}).get("pdf2image")) or has_python_module("pdf2image"),
            "pdfplumber": bool(docs_probe.get("modules", {}).get("pdfplumber")) or has_python_module("pdfplumber"),
            "pypdf": bool(docs_probe.get("modules", {}).get("pypdf")) or has_python_module("pypdf"),
            "reportlab": bool(docs_probe.get("modules", {}).get("reportlab")) or has_python_module("reportlab"),
            "pypdfium2": bool(docs_probe.get("modules", {}).get("pypdfium2")) or has_python_module("pypdfium2"),
            "playwright": has_python_module("playwright"),
        },
        "codex_mcp": mcp,
        "notion_mcp_configured": notion_configured,
        "notion_credentials_present": bool(core_probe.get("notion_credentials_present")),
        "rmcp_enabled": rmcp_enabled,
        "runtime_root": str(resolve_runtime_root(runtime_root)),
        "runtime_registry_path": str(registry_path(runtime_root)),
        "runtime_registry": runtime_registry,
    }


def readiness_for_skill(skill_name: str, skill_path: Path, env: dict[str, Any]) -> dict[str, Any]:
    valid, message = validate_skill(skill_path)
    checks: dict[str, Any] = {"skill_valid": valid, "validation_message": message}
    docs_modules = env.get("python_modules", {})
    core_probe = env.get("runtime_registry", {}).get("domains", {}).get("core-tools", {}).get("probe", {})
    docs_probe = env.get("runtime_registry", {}).get("domains", {}).get("docs-python", {}).get("probe", {})

    if skill_name == "doc":
        checks.update(
            {
                "runtime_domain": "docs-python",
                "runtime_root": env.get("runtime_root"),
                "runtime_python": docs_probe.get("python"),
                "runtime_available": docs_probe.get("available", False),
                "python_docx": docs_modules.get("docx", False),
                "pdf2image": docs_modules.get("pdf2image", False),
                "pypdfium2": docs_modules.get("pypdfium2", False),
                "soffice_present": core_probe.get("libreoffice_present", False),
            }
        )
        if docs_probe.get("available") and docs_modules.get("docx") and docs_modules.get("pdf2image") and docs_modules.get("pypdfium2") and core_probe.get("libreoffice_present"):
            verdict = "ready"
            reason = "Shared docs runtime and LibreOffice are available for DOCX rendering"
        elif docs_probe.get("python"):
            verdict = "conditional_ready"
            reason = "Shared docs runtime exists, but DOCX visual-render dependencies are still incomplete"
        else:
            verdict = "env_blocked"
            reason = "Shared docs runtime is not bootstrapped yet"
    elif skill_name == "pdf":
        checks.update(
            {
                "runtime_domain": "docs-python",
                "runtime_root": env.get("runtime_root"),
                "runtime_python": docs_probe.get("python"),
                "runtime_available": docs_probe.get("available", False),
                "pdfplumber": docs_modules.get("pdfplumber", False),
                "pypdf": docs_modules.get("pypdf", False),
                "reportlab": docs_modules.get("reportlab", False),
                "pypdfium2": docs_modules.get("pypdfium2", False),
            }
        )
        if docs_probe.get("available") and docs_modules.get("pdfplumber") and docs_modules.get("pypdf") and docs_modules.get("reportlab") and docs_modules.get("pypdfium2"):
            verdict = "ready"
            reason = "Shared docs runtime is available for PDF generation, extraction, and rendering"
        elif docs_probe.get("python"):
            verdict = "conditional_ready"
            reason = "Shared docs runtime exists, but PDF helper libraries are not fully installed"
        else:
            verdict = "env_blocked"
            reason = "Shared docs runtime is not bootstrapped yet"
    elif skill_name == "gh-fix-ci":
        checks.update(
            {
                "gh_present": env["gh_present"],
                "gh_authenticated": env["gh_authenticated"],
                "gh_auth_message": env["gh_auth_message"],
                "git_present": env["git_present"],
            }
        )
        if env["gh_present"] and env["gh_authenticated"]:
            verdict = "ready"
            reason = "GitHub CLI is installed and authenticated for read-only CI inspection workflows"
        elif env["gh_present"]:
            verdict = "conditional_ready"
            reason = "GitHub CLI is installed, but authentication is still required"
        else:
            verdict = "env_blocked"
            reason = "GitHub CLI is not installed in this environment"
    elif skill_name == "yeet":
        checks.update(
            {
                "gh_present": env["gh_present"],
                "gh_authenticated": env["gh_authenticated"],
                "gh_auth_message": env["gh_auth_message"],
                "git_present": env["git_present"],
            }
        )
        if env["gh_present"] and env["gh_authenticated"]:
            verdict = "ready"
            reason = "GitHub CLI is installed and authenticated; external writes still require explicit user confirmation"
        elif env["gh_present"]:
            verdict = "conditional_ready"
            reason = "GitHub CLI is installed, but authenticated user intent is still required"
        else:
            verdict = "env_blocked"
            reason = "GitHub CLI is not installed in this environment"
    elif skill_name in NOTION_SKILLS:
        checks.update(
            {
                "codex_mcp_configured": env["codex_mcp"]["configured"],
                "notion_mcp_configured": env["notion_mcp_configured"],
                "rmcp_enabled": env["rmcp_enabled"],
                "notion_credentials_present": env["notion_credentials_present"],
            }
        )
        if env["notion_mcp_configured"] and env["rmcp_enabled"] and env["notion_credentials_present"]:
            verdict = "ready"
            reason = "Notion MCP is configured and OAuth credentials are present"
        elif env["notion_mcp_configured"]:
            verdict = "conditional_ready"
            reason = "Notion MCP is configured, but OAuth authentication is still required"
        else:
            verdict = "env_blocked"
            reason = "Notion MCP is not configured in this environment"
    elif skill_name == "playwright":
        checks.update(
            {
                "playwright_cli_present": env["playwright_present"],
                "python_playwright": env["python_modules"]["playwright"],
                "playwright_mcp_configured": env["playwright_mcp_configured"],
                "playwright_browser_cache_present": env["playwright_browser_cache_present"],
            }
        )
        if env["playwright_present"] and env["playwright_browser_cache_present"]:
            verdict = "ready"
            reason = "Playwright CLI and browser runtime are installed locally"
        elif env["playwright_mcp_configured"]:
            verdict = "conditional_ready"
            reason = "Playwright MCP is configured, but the local CLI/browser runtime is incomplete"
        else:
            verdict = "env_blocked"
            reason = "Playwright runtime is not installed in this environment"
    elif skill_name == "screenshot":
        checks.update(
            {
                "display": env["display"],
                "wayland_display": env["wayland_display"],
                **env["linux_screenshot_tools"],
            }
        )
        if any(bool(value) for value in env["linux_screenshot_tools"].values()):
            verdict = "ready"
            reason = "Linux screenshot backend available"
        else:
            verdict = "env_blocked"
            reason = "No supported Linux screenshot utility is installed"
    elif skill_name == "jupyter-notebook":
        checks.update(
            {
                "python3_present": env["python3_present"],
                "script_present": (skill_path / "scripts" / "new_notebook.py").exists(),
                "template_count": len(list((skill_path / "assets").glob("*.ipynb"))),
            }
        )
        verdict = "ready"
        reason = "Python, notebook scaffolding script, and templates are available locally"
    elif skill_name == "latex-to-word":
        checks.update({"pandoc_present": env["pandoc_present"], "runtime_domain": "docs-python"})
        verdict = "ready" if env["pandoc_present"] else "env_blocked"
        reason = "Pandoc is installed locally" if env["pandoc_present"] else "Pandoc is missing"
    elif skill_name == "openai-docs":
        verdict = "ready"
        reason = "This skill is guidance-first and does not require extra local runtimes"
    else:
        verdict = "ready" if valid else "env_blocked"
        reason = "Default readiness applied" if valid else f"skill validation failed: {message}"

    if not valid:
        verdict = "env_blocked"
        reason = f"skill validation failed: {message}"

    return {
        "readiness_verdict": verdict,
        "readiness_reason": reason,
        "checks": checks,
    }


def load_bundled_evaluations(skill_path: Path, skill_name: str) -> dict[str, Any] | None:
    eval_dir = skill_path / "evaluations"
    if not eval_dir.is_dir():
        return None
    evals = []
    for index, path in enumerate(sorted(eval_dir.glob("*.json")), start=1):
        payload = load_json(path)
        prompt = payload.get("query") or payload.get("prompt") or f"Use ${skill_name} on a representative task from {path.stem}."
        expected_bits = list(payload.get("success_criteria", [])) + list(payload.get("expected_behavior", []))
        evals.append(
            {
                "id": index,
                "name": payload.get("name", path.stem),
                "mode": "task",
                "files": [str(path.relative_to(eval_dir.parent))],
                "prompt": prompt,
                "expected_output": " ".join(expected_bits[:8]) or payload.get("name", path.stem),
                "workflow_expectations": {
                    "required_steps": list(payload.get("expected_behavior", []))[:8],
                    "required_artifacts": [],
                    "critical_constraints": list(payload.get("success_criteria", []))[:8],
                },
            }
        )
    return {
        "skill_name": skill_name,
        "status": "bundled_draft",
        "executable": False,
        "notes": [
            "Converted from bundled skill evaluations for audit context; not auto-benchmarked in this run.",
        ],
        "evals": evals,
    }


def write_fixture_file(path: Path, content: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    dump_text(path, content)


def _compile_eval_pack(pack_dir: Path, payload: dict[str, Any]) -> dict[str, Any]:
    compiled = {
        "skill_name": payload["skill_name"],
        "status": payload.get("status", "auto_generated"),
        "executable": bool(payload.get("executable", False)),
        "notes": payload.get("notes", []),
        "evals": payload.get("evals", []),
        "source_spec": payload,
    }
    dump_json(pack_dir / "simple_eval_spec.json", payload)
    dump_json(pack_dir / "compiled_evals.json", compiled)
    return compiled


def build_screenshot_eval_pack(*, skill_name: str, pack_dir: Path) -> dict[str, Any]:
    fixtures_dir = pack_dir / "fixtures"
    write_fixture_file(
        fixtures_dir / "start_test_window.sh",
        "#!/usr/bin/env bash\n"
        "set -euo pipefail\n"
        "title=\"${1:-Codex Screenshot Audit}\"\n"
        "body=\"${2:-Codex Screenshot Audit Window}\"\n"
        "timeout 45 xmessage -center -title \"$title\" \"$body\" >/dev/null 2>&1 &\n"
        "echo $! > xmessage.pid\n"
        "sleep 2\n",
    )
    (fixtures_dir / "start_test_window.sh").chmod(0o755)
    payload = {
        "skill_name": skill_name,
        "status": "auto_generated",
        "executable": True,
        "generated_by": "skill_manager_local_eval_pack",
        "notes": [
            "Uses a controlled xmessage window so the audit does not capture unrelated desktop content.",
        ],
        "evals": [
            {
                "id": 1,
                "name": "capture-controlled-active-window",
                "mode": "task",
                "files": ["fixtures/start_test_window.sh"],
                "prompt": (
                    f"Use ${skill_name} to capture a controlled test window instead of the real desktop. "
                    "Run `bash fixtures/start_test_window.sh 'Codex Screenshot Audit' 'Codex Screenshot Audit Window'`, "
                    "save the capture to `captured-active.png`, and record the saved path in `capture-note.md`."
                ),
                "expected_output": "Creates a screenshot file for the controlled test window and reports the saved path.",
                "must_create": ["captured-active.png", "capture-note.md"],
                "must_include": ["captured-active.png"],
                "workflow_expectations": {
                    "required_steps": [
                        "Create the controlled test window instead of capturing the user's real desktop.",
                        "Capture the active window only.",
                        "Report the saved screenshot path.",
                    ],
                    "required_artifacts": ["captured-active.png", "capture-note.md"],
                    "critical_constraints": [
                        "Do not capture unrelated desktop content.",
                    ],
                },
            }
        ],
    }
    return _compile_eval_pack(pack_dir, payload)


def build_gh_fix_ci_eval_pack(*, skill_name: str, pack_dir: Path) -> dict[str, Any]:
    fixtures_dir = pack_dir / "fixtures"
    write_fixture_file(fixtures_dir / "job_987654321.log", "E   ModuleNotFoundError: No module named 'yaml'\n")
    payload = {
        "skill_name": skill_name,
        "status": "auto_generated",
        "executable": True,
        "generated_by": "skill_manager_local_eval_pack",
        "notes": [
            "Uses local fixtures so the benchmark stays read-only and deterministic.",
        ],
        "evals": [
            {
                "id": 1,
                "name": "triage-github-actions-fixture",
                "mode": "task",
                "files": ["fixtures/job_987654321.log"],
                "prompt": (
                    f"Use ${skill_name} to summarize a failing GitHub Actions check from local fixtures only, "
                    "then create `fix-plan.md` with the failure summary and an approval-gated fix plan."
                ),
                "expected_output": "Creates a triage file with the failing check summary and a fix plan.",
                "must_create": ["fix-plan.md"],
                "must_include": ["fix-plan.md"],
                "workflow_expectations": {
                    "required_steps": [
                        "Inspect the provided local fixture.",
                        "Summarize the failing check.",
                        "Draft a fix plan without implementing changes.",
                    ],
                    "required_artifacts": ["fix-plan.md"],
                    "critical_constraints": [
                        "Do not implement code changes before explicit approval.",
                    ],
                },
            }
        ],
    }
    return _compile_eval_pack(pack_dir, payload)


def build_doc_eval_pack(*, skill_name: str, pack_dir: Path) -> dict[str, Any]:
    fixtures_dir = pack_dir / "fixtures"
    write_fixture_file(
        fixtures_dir / "create_sample_docx.py",
        "from pathlib import Path\n"
        "import sys\n"
        "from docx import Document\n\n"
        "target = Path(sys.argv[1])\n"
        "target.parent.mkdir(parents=True, exist_ok=True)\n"
        "doc = Document()\n"
        "doc.add_heading('Codex DOCX Audit', level=1)\n"
        "doc.add_paragraph('This document exists only for the controlled skill audit.')\n"
        "table = doc.add_table(rows=2, cols=2)\n"
        "table.cell(0, 0).text = 'Column A'\n"
        "table.cell(0, 1).text = 'Column B'\n"
        "table.cell(1, 0).text = 'Alpha'\n"
        "table.cell(1, 1).text = 'Beta'\n"
        "doc.save(target)\n",
    )
    write_fixture_file(
        fixtures_dir / "update_sample_docx.py",
        "from pathlib import Path\n"
        "import sys\n"
        "from docx import Document\n\n"
        "target = Path(sys.argv[1])\n"
        "doc = Document(target)\n"
        "doc.add_paragraph('Boundary update for the controlled audit document.')\n"
        "doc.save(target)\n",
    )
    for script_name in ("create_sample_docx.py", "update_sample_docx.py"):
        (fixtures_dir / script_name).chmod(0o755)
    payload = {
        "skill_name": skill_name,
        "status": "auto_generated",
        "executable": True,
        "generated_by": "skill_manager_local_eval_pack",
        "notes": [
            "Uses a controlled DOCX fixture and the bundled render helper through the shared docs runtime.",
        ],
        "evals": [
            {
                "id": 1,
                "name": "render-controlled-docx",
                "mode": "task",
                "files": ["fixtures/create_sample_docx.py"],
                "prompt": (
                    f"Use ${skill_name} to complete a controlled DOCX render handoff. "
                    "Run `.agents/skills/doc/scripts/use_docs_runtime.sh fixtures/create_sample_docx.py sample.docx`, "
                    "then render it with `.agents/skills/doc/scripts/use_docs_runtime.sh .agents/skills/doc/scripts/render_docx.py sample.docx --output_dir rendered-docx`, "
                    "and create `doc-review.md` with the saved path and render backend."
                ),
                "expected_output": "Creates a controlled DOCX fixture, renders it with the bundled helper, and writes a deterministic review note.",
                "must_create": ["sample.docx", "rendered-docx/page-1.png", "rendered-docx/render-manifest.json", "doc-review.md"],
                "must_include": ["rendered-docx/page-1.png"],
                "workflow_expectations": {
                    "required_steps": [
                        "Use the shared docs runtime launcher.",
                        "Create a controlled DOCX fixture.",
                        "Render the DOCX with the bundled helper.",
                        "Record the render backend and deterministic review checkpoints.",
                    ],
                    "required_artifacts": ["sample.docx", "rendered-docx/page-1.png", "rendered-docx/render-manifest.json", "doc-review.md"],
                    "critical_constraints": [
                        "Do not modify files under `.agents/skills/doc/`.",
                        "Do not replace the bundled render helper with an HTML-only shortcut.",
                        "If render_backend is structured-fallback, do not claim full LibreOffice layout fidelity.",
                    ],
                },
            },
            {
                "id": 2,
                "name": "doc-boundary-and-rerender",
                "mode": "task",
                "files": ["fixtures/create_sample_docx.py", "fixtures/update_sample_docx.py"],
                "prompt": (
                    f"Use ${skill_name} on a controlled DOCX task that also requests one unsupported extra. "
                    "Create and update the sample DOCX, rerender it to `rerendered-docx`, and create `doc-boundary.md` "
                    "that records the render backend plus the unsupported boundary."
                ),
                "expected_output": "Rerenders the updated DOCX and clearly refuses the unsupported extra.",
                "must_create": ["sample.docx", "rerendered-docx/page-1.png", "rerendered-docx/render-manifest.json", "doc-boundary.md"],
                "must_include": ["rerendered-docx/page-1.png"],
                "workflow_expectations": {
                    "required_steps": [
                        "Create and update the controlled DOCX fixture through the shared docs runtime.",
                        "Rerender the DOCX after the edit with the bundled helper.",
                        "State the unsupported extra explicitly instead of pretending the skill covers it.",
                        "Record the rerender backend in the boundary note.",
                    ],
                    "required_artifacts": ["sample.docx", "rerendered-docx/page-1.png", "rerendered-docx/render-manifest.json", "doc-boundary.md"],
                    "critical_constraints": [
                        "Do not fabricate support for the unsupported extra.",
                        "Do not modify files under `.agents/skills/doc/`.",
                        "Do not replace the bundled render helper with an HTML-only shortcut.",
                        "If render_backend is structured-fallback, do not claim full LibreOffice layout fidelity.",
                    ],
                },
            },
        ],
    }
    return _compile_eval_pack(pack_dir, payload)


def build_pdf_eval_pack(*, skill_name: str, pack_dir: Path) -> dict[str, Any]:
    fixtures_dir = pack_dir / "fixtures"
    write_fixture_file(
        fixtures_dir / "create_sample_pdf.py",
        "from pathlib import Path\n"
        "import sys\n"
        "from reportlab.lib.pagesizes import letter\n"
        "from reportlab.pdfgen import canvas\n\n"
        "target = Path(sys.argv[1])\n"
        "target.parent.mkdir(parents=True, exist_ok=True)\n"
        "pdf = canvas.Canvas(str(target), pagesize=letter)\n"
        "pdf.setTitle('Codex PDF Audit')\n"
        "pdf.drawString(72, 720, 'Codex PDF Audit Document')\n"
        "pdf.drawString(72, 690, 'This PDF exists only for the controlled skill audit.')\n"
        "pdf.save()\n",
    )
    (fixtures_dir / "create_sample_pdf.py").chmod(0o755)
    payload = {
        "skill_name": skill_name,
        "status": "auto_generated",
        "executable": True,
        "generated_by": "skill_manager_local_eval_pack",
        "notes": [
            "Uses a controlled PDF fixture and the bundled render helper through the shared docs runtime.",
        ],
        "evals": [
            {
                "id": 1,
                "name": "render-controlled-pdf",
                "mode": "task",
                "files": ["fixtures/create_sample_pdf.py"],
                "prompt": (
                    f"Use ${skill_name} to create and visually verify a controlled PDF fixture. "
                    "Render the PDF to `rendered-pdf/page-1.png`, then create `pdf-review.md` with one concrete visual check."
                ),
                "expected_output": "Creates a controlled PDF fixture, renders it visually, and records a review note.",
                "must_create": ["sample.pdf", "rendered-pdf/page-1.png", "pdf-review.md"],
                "must_include": ["rendered-pdf/page-1.png"],
                "workflow_expectations": {
                    "required_steps": [
                        "Use the shared docs runtime launcher.",
                        "Create a controlled PDF fixture.",
                        "Render the PDF visually with the bundled helper.",
                        "Record at least one visual review conclusion.",
                    ],
                    "required_artifacts": ["sample.pdf", "rendered-pdf/page-1.png", "pdf-review.md"],
                    "critical_constraints": [
                        "Do not claim visual review without rendering the PDF.",
                    ],
                },
            }
        ],
    }
    return _compile_eval_pack(pack_dir, payload)


def ensure_eval_pack_for_skill(
    skill_entry: dict[str, Any],
    workspace_root: Path,
    *,
    use_bundled_if_available: bool = True,
) -> dict[str, Any]:
    skill_name = skill_entry["skill_name"]
    skill_path = Path(skill_entry["skill_path"])
    pack_dir = workspace_root / skill_name
    pack_dir.mkdir(parents=True, exist_ok=True)

    if use_bundled_if_available:
        bundled = load_bundled_evaluations(skill_path, skill_name)
        if bundled:
            dump_json(pack_dir / "bundled_evaluations_simple_eval_spec.json", bundled)
            return _compile_eval_pack(pack_dir, bundled)

    builders = {
        "doc": build_doc_eval_pack,
        "pdf": build_pdf_eval_pack,
        "screenshot": build_screenshot_eval_pack,
        "gh-fix-ci": build_gh_fix_ci_eval_pack,
    }
    if skill_name in builders:
        return builders[skill_name](skill_name=skill_name, pack_dir=pack_dir)

    draft = {
        "skill_name": skill_name,
        "status": "manual_follow_up",
        "executable": False,
        "notes": [
            "No bundled or auto-generated safe eval pack is available for this skill.",
        ],
        "evals": [],
    }
    return _compile_eval_pack(pack_dir, draft)


def write_markdown_report(payload: dict[str, Any], output_path: Path) -> None:
    lines = [
        "# Migrated Skills Audit",
        "",
        f"- Generated at: {payload['generated_at']}",
        f"- Manifest: `{payload['manifest_path']}`",
        f"- Skills root: `{payload['skills_root']}`",
        "",
        "| Skill | Readiness | Retirement | Reason |",
        "| --- | --- | --- | --- |",
    ]
    for item in payload["skills"]:
        lines.append(
            f"| `{item['skill_name']}` | `{item['readiness_verdict']}` | `{item['retirement_verdict']}` | {item['retirement_reason']} |"
        )
    dump_text(output_path, "\n".join(lines) + "\n")


def audit_skills(
    *,
    manifest_path: Path,
    workspace_root: Path,
    skills_root: Path = DEFAULT_SKILLS_ROOT,
    rerun_skills: set[str] | None = None,
    runtime_root: str | Path | None = None,
) -> dict[str, Any]:
    layout = ensure_lab_workspace(workspace_root)
    manifest = load_json(manifest_path)
    env = probe_environment(runtime_root=runtime_root)
    inventory = discover_skills(skills_root, layout["root"])
    dump_json(
        layout["root"] / "post_migration_catalog.json",
        {
            "generated_at": utc_now_iso(),
            "skills_root": str(Path(skills_root).expanduser().resolve()),
            "skills": inventory,
        },
    )

    rerun_targets = set(rerun_skills or set())
    copied = manifest.get("copied_skills") or []
    skills: list[dict[str, Any]] = []
    for item in copied:
        skill_name = str(item.get("skill_name") or "").strip()
        if not skill_name:
            continue
        skill_path = Path(item.get("dest_path") or (Path(skills_root) / skill_name)).expanduser()
        readiness = readiness_for_skill(skill_name, skill_path, env)
        record = {
            "skill_name": skill_name,
            "skill_path": str(skill_path),
            "source_path": item.get("source_path"),
            "readiness_verdict": readiness["readiness_verdict"],
            "readiness_reason": readiness["readiness_reason"],
            "checks": readiness["checks"],
        }

        if readiness["readiness_verdict"] == "env_blocked":
            record["retirement_verdict"] = "env_blocked"
            record["retirement_reason"] = readiness["readiness_reason"]
        elif skill_name in READY_EXTERNAL_WRITE_SKILLS:
            record["retirement_verdict"] = "keep_active"
            record["retirement_reason"] = "safe_skip_external_side_effects"
        else:
            use_auto_pack = skill_name in SAFE_BENCHMARK_SKILLS or skill_name in rerun_targets
            eval_pack = ensure_eval_pack_for_skill(
                record,
                layout["eval_packs"],
                use_bundled_if_available=not use_auto_pack,
            )
            record["eval_pack_executable"] = eval_pack["executable"]
            record["eval_pack_path"] = str(layout["eval_packs"] / skill_name / "compiled_evals.json")
            if eval_pack["executable"]:
                record["retirement_verdict"] = "keep_active"
                record["retirement_reason"] = "safe_local_eval_pack"
            else:
                record["retirement_verdict"] = "needs_manual_review"
                record["retirement_reason"] = "no_safe_local_eval_pack"

        skills.append(record)

    readiness_payload = {
        "generated_at": utc_now_iso(),
        "manifest_path": str(manifest_path.resolve()),
        "skills_root": str(Path(skills_root).expanduser().resolve()),
        "runtime_root": env["runtime_root"],
        "runtime_registry_path": env["runtime_registry_path"],
        "skills": skills,
    }
    dump_json(layout["root"] / "readiness_report.json", readiness_payload)

    retirement_payload = {
        "generated_at": utc_now_iso(),
        "manifest_path": str(manifest_path.resolve()),
        "skills_root": str(Path(skills_root).expanduser().resolve()),
        "skills": skills,
    }
    dump_json(layout["reports"] / "retirement_scan.json", retirement_payload)
    write_markdown_report(retirement_payload, layout["reports"] / "retirement_scan.md")
    return {
        "readiness_report": readiness_payload,
        "retirement_scan": retirement_payload,
    }


def main(argv: list[str]) -> int:
    parser = argparse.ArgumentParser(description="Audit migrated skills using only local readiness and safe eval-pack checks")
    parser.add_argument("--manifest", required=True, help="Path to migration_manifest.json")
    parser.add_argument("--workspace", required=True, help="Workspace for reports and generated eval packs")
    parser.add_argument("--skills-root", default=str(DEFAULT_SKILLS_ROOT))
    parser.add_argument("--runtime-root", default=None)
    parser.add_argument("--rerun-skill", action="append", default=[], help="Force auto eval-pack generation for the given skill")
    args = parser.parse_args(argv)

    result = audit_skills(
        manifest_path=Path(args.manifest).expanduser().resolve(),
        workspace_root=Path(args.workspace).expanduser().resolve(),
        skills_root=Path(args.skills_root).expanduser(),
        rerun_skills={item for item in args.rerun_skill if item},
        runtime_root=args.runtime_root,
    )
    print(json.dumps(result["retirement_scan"]["skills"], indent=2, ensure_ascii=False))
    return 0


if __name__ == "__main__":
    raise SystemExit(main(sys.argv[1:]))
